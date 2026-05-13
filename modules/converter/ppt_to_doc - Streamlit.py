import os
import tempfile
import subprocess

from docx import Document
from docx.shared import Inches

from modules.converter.ppt_extractor import extract_ppt_content
from modules.converter.ppt_slide_renderer import render_ppt_slides_to_images


def convert_clean_ppt_to_images(clean_ppt_path):
    """
    Convert cleaned PPT into PNG images
    """
    try:
        output_dir = tempfile.mkdtemp()

        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to",
            "png",
            clean_ppt_path,
            "--outdir",
            output_dir
        ], check=True)

        images = []

        for file in sorted(os.listdir(output_dir)):
            if file.lower().endswith(".png"):
                images.append(
                    os.path.join(output_dir, file)
                )

        print(f"Converted images: {images}")
        return images

    except Exception as e:
        print(f"PNG conversion failed: {e}")
        return []


def add_images_to_doc(doc, image_paths):
    """
    Add images to word
    """
    for img_path in image_paths:
        try:
            if os.path.exists(img_path):
                doc.add_picture(
                    img_path,
                    width=Inches(6.5)
                )
                doc.add_page_break()

        except Exception as e:
            print(f"Image insert failed: {e}")


def convert_ppt_to_doc(ppt_path, output_docx):
    try:
        doc = Document()

        # ---------------------------
        # Extract text
        # ---------------------------
        extracted_text = extract_ppt_content(ppt_path)

        if extracted_text:
            doc.add_heading("PPT Content", level=1)

            for text in extracted_text:
                if text.strip():
                    doc.add_paragraph(text)

        # ---------------------------
        # Clean ppt
        # ---------------------------
        cleaned_ppt_path = render_ppt_slides_to_images(
            ppt_path
        )

        print(f"Clean PPT path: {cleaned_ppt_path}")

        # ---------------------------
        # Convert cleaned ppt → png
        # ---------------------------
        image_paths = convert_clean_ppt_to_images(
            cleaned_ppt_path
        )

        print(f"Final images: {image_paths}")

        if image_paths:
            doc.add_page_break()
            doc.add_heading("PPT Slides", level=1)

            add_images_to_doc(
                doc,
                image_paths
            )
        else:
            doc.add_paragraph(
                "No PPT slide images generated."
            )

        doc.save(output_docx)

        return output_docx

    except Exception as e:
        print(f"PPT conversion failed: {e}")
        raise


def ppt_to_word(ppt_path, output_docx):
    return convert_ppt_to_doc(
        ppt_path,
        output_docx
    )
