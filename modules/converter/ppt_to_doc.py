import os

from docx import Document
from docx.shared import Inches

from modules.converter.ppt_extractor import extract_ppt_content
from modules.converter.ppt_slide_renderer import render_ppt_slides_to_images
from modules.common.logger import setup_logger


logger = setup_logger("ppt_to_doc")


def add_images_to_doc(doc, image_paths):
    for img_path in image_paths:
        try:
            if os.path.exists(img_path):
                logger.info(
                    f"Adding image to doc: {img_path}"
                )

                doc.add_picture(
                    img_path,
                    width=Inches(6.5)
                )

                doc.add_page_break()

            else:
                logger.warning(
                    f"Image missing: {img_path}"
                )

        except Exception as e:
            logger.error(
                f"Image insert failed: {str(e)}"
            )


def convert_ppt_to_doc(ppt_path, output_docx):
    try:
        logger.info(
            f"Starting PPT->DOC conversion: {ppt_path}"
        )

        doc = Document()

        extracted_text = extract_ppt_content(
            ppt_path
        )

        logger.info(
            f"Extracted text count: {len(extracted_text)}"
        )

        if extracted_text:
            doc.add_heading(
                "PPT Content",
                level=1
            )

            for text in extracted_text:
                if text.strip():
                    doc.add_paragraph(text)

        image_paths = render_ppt_slides_to_images(
            ppt_path
        )

        logger.info(
            f"Rendered image count: {len(image_paths)}"
        )

        if image_paths:
            doc.add_page_break()
            doc.add_heading(
                "PPT Slides",
                level=1
            )

            add_images_to_doc(
                doc,
                image_paths
            )
        else:
            logger.warning(
                "No slide images generated"
            )

            doc.add_paragraph(
                "No PPT slide images generated."
            )

        doc.save(output_docx)

        logger.info(
            f"DOC saved successfully: {output_docx}"
        )

        return output_docx

    except Exception as e:
        logger.error(
            f"PPT conversion failed: {str(e)}"
        )
        raise


def ppt_to_word(ppt_path, output_docx):
    logger.info("ppt_to_word called")

    return convert_ppt_to_doc(
        ppt_path,
        output_docx
    )