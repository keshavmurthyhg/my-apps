import os

from docx import Document
from docx.shared import Inches

from modules.converter.ppt_extractor import extract_ppt_content
from modules.converter.ppt_slide_renderer import (
    render_ppt_slides_to_images
)


# ----------------------------------------
# Add extracted images into Word
# ----------------------------------------
def add_images_to_doc(doc, image_paths):
    """
    Insert extracted slide images
    into Word document
    """
    for index, img_path in enumerate(
        image_paths,
        start=1
    ):
        try:
            if not os.path.exists(img_path):
                print(
                    f"Image not found: {img_path}"
                )
                continue

            doc.add_heading(
                f"Slide {index}",
                level=2
            )

            doc.add_picture(
                img_path,
                width=Inches(6.5)
            )

            doc.add_page_break()

            print(
                f"Inserted image: {img_path}"
            )

        except Exception as e:
            print(
                f"Image insert failed "
                f"for {img_path}: {e}"
            )


# ----------------------------------------
# Convert PPT → Word
# ----------------------------------------
def convert_ppt_to_doc(
    ppt_path,
    output_docx
):
    try:
        doc = Document()

        # --------------------------------
        # Extract textual content
        # --------------------------------
        extracted_text = extract_ppt_content(
            ppt_path
        )

        if extracted_text:
            doc.add_heading(
                "PPT Content",
                level=1
            )

            for text in extracted_text:
                if text.strip():
                    doc.add_paragraph(text)

        else:
            print(
                "No text extracted from PPT"
            )

        # --------------------------------
        # Extract slide images directly
        # --------------------------------
        image_paths = render_ppt_slides_to_images(
            ppt_path
        )

        print(
            f"Final extracted images: "
            f"{image_paths}"
        )

        # --------------------------------
        # Add images to document
        # --------------------------------
        if image_paths:
            doc.add_page_break()

            doc.add_heading(
                "PPT Slides",
                level=1
            )

            add_images_to_doc(
                doc,
                image_paths
            )

        else:
            print(
                "No slide images found"
            )

            doc.add_page_break()
            doc.add_heading(
                "PPT Slides",
                level=1
            )

            doc.add_paragraph(
                "No slide images found "
                "in PPT."
            )

        # --------------------------------
        # Save final doc
        # --------------------------------
        doc.save(output_docx)

        print(
            f"Word document created: "
            f"{output_docx}"
        )

        return output_docx

    except Exception as e:
        print(
            f"PPT conversion failed: {e}"
        )
        raise


# ----------------------------------------
# Wrapper
# ----------------------------------------
def ppt_to_word(
    ppt_path,
    output_docx
):
    return convert_ppt_to_doc(
        ppt_path,
        output_docx
    )